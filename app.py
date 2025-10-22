import streamlit as st
from PIL import Image
import io, requests, base64, os, tempfile
from docx import Document
from docx.shared import Inches

st.set_page_config(page_title="Auto Rubric TK Pertiwi", layout="wide")

# Pink pastel style via simple header
st.markdown("<style>body{background: linear-gradient(180deg,#fff0f6,#fffdf8);} .stApp {font-family: Inter, Arial, sans-serif;}</style>", unsafe_allow_html=True)

col1, col2 = st.columns([1,5])
with col1:
    st.image("logo.png", width=96)
with col2:
    st.title("Laporan Hasil Karya Anak - TK Pertiwi")
    st.markdown("TAMAN KANAK–KANAK PERTIWI 26–68 — PRUPUK UTARA, KECAMATAN MARGASARI, KABUPATEN TEGAL")

st.markdown("---")

# Identity block (prefilled)
with st.expander("Identitas (dapat diubah)"):
    cols = st.columns(3)
    kelompok = cols[0].text_input("Kelompok", value="B5")
    tanggal = cols[1].text_input("Hari / Tanggal", value="Rabu, 23 Juli 2025")
    semester = cols[2].text_input("Semester", value="I")
    cols2 = st.columns(3)
    teacher = cols2[0].text_input("Guru Kelas (opsional)", value="Kurnia Khoerul Aniza, S.Pd")
    topic = cols2[1].text_input("Topik", value="Aku dan Lingkunganku")
    subtopic = cols2[2].text_input("Sub Topik", value="Aku Sayang Tubuhku")

st.info("Upload tepat 3 foto per laporan (urut sesuai baris nama). AI akan membuat 4 poin deskripsi per foto. Untuk publik app tanpa memasukkan token setiap pengguna, set HUGGINGFACE_TOKEN in Streamlit Secrets (hanya sekali).")

# Admin panel - session token (optional)
with st.expander("Admin — Token AI (opsional, sekali per sesi)"):
    if "hf_token" not in st.session_state:
        st.session_state["hf_token"] = ""
    token_input = st.text_input("Hugging Face token (hf_...)", value=st.session_state["hf_token"], type="password")
    if st.button("Simpan token sesi"):
        st.session_state["hf_token"] = token_input.strip()
        st.success("Token disimpan di sesi. Untuk permanen, set secret HUGGINGFACE_TOKEN pada hosting (Streamlit Cloud).")

# Upload exactly 3 images
uploaded = st.file_uploader("Upload 3 foto karya (pilih 3 file)", type=['png','jpg','jpeg'], accept_multiple_files=True)
if uploaded and len(uploaded) != 3:
    st.warning("Mohon upload tepat 3 file. Saat ini: %d file terpilih." % len(uploaded))

if uploaded and len(uploaded) == 3:
    for idx, file in enumerate(uploaded):
        cols = st.columns([1,2])
        with cols[0]:
            name = st.text_input(f"Nama anak baris {idx+1}", key=f"name_{idx}")
            img = Image.open(file).convert("RGB")
            st.image(img, width=240)
        with cols[1]:
            st.markdown("**Deskripsi (AI-generated, bullet)**")
            if st.button(f"Analisis AI baris {idx+1}", key=f"ai_{idx}"):
                prompt = f"Konteks: hasil karya anak usia taman kanak-kanak yang dikerjakan secara mandiri. Nama anak: {name}. Buat 4 poin singkat (bullet) bergaya naratif formal PAUD tentang apa yang terlihat pada foto karya (pemilihan bahan/warna, teknik: menempel/robek/menggambar, kerapian, kreativitas). JANGAN ulangi nama anak dalam kalimat; jika model menyebut nama, gunakan kata 'anak'. Tulis setiap poin di baris baru."
                hf_token = os.getenv('HUGGINGFACE_TOKEN') or st.session_state.get('hf_token','')
                if not hf_token:
                    st.warning("Token tidak tersedia — menggunakan deskripsi fallback.")
                    desc = "• Memilih media dan bahan secara mandiri\n• Menyusun dan menempel dengan rapi sesuai ide\n• Menunjukkan kreativitas dan percaya diri saat berkarya\n• Menghasilkan karya yang menarik dan bermakna"
                else:
                    try:
                        headers = {"Authorization": f"Bearer {hf_token}", "Content-Type": "application/json"}
                        model_url = "https://api-inference.huggingface.co/models/google/flan-t5-small"
                        resp = requests.post(model_url, headers=headers, json={"inputs": prompt, "parameters": {"max_new_tokens": 200}})
                        resp.raise_for_status()
                        data = resp.json()
                        text = ""
                        if isinstance(data, list) and len(data)>0 and "generated_text" in data[0]:
                            text = data[0]["generated_text"]
                        elif isinstance(data, dict) and "generated_text" in data:
                            text = data["generated_text"]
                        elif isinstance(data, str):
                            text = data
                        else:
                            text = str(data)
                        if name:
                            text = text.replace(name, "anak")
                        lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
                        if not lines:
                            parts = [p.strip() for p in text.split(".") if p.strip()]
                            lines = parts[:4]
                        lines = lines[:4]
                        desc = "\n".join(["• "+ln for ln in lines])
                    except Exception as e:
                        st.error("Error memanggil AI: "+str(e))
                        desc = "• Memilih media dan bahan secara mandiri\n• Menyusun dan menempel dengan rapi sesuai ide\n• Menunjukkan kreativitas dan percaya diri saat berkarya\n• Menghasilkan karya yang menarik dan bermakna"
                st.session_state[f"desc_{idx}"] = desc
            desc_area = st.text_area(f"Deskripsi baris {idx+1}", value=st.session_state.get(f"desc_{idx}",""), height=160, key=f"desc_area_{idx}")
            if st.button(f"Simpan baris {idx+1}", key=f"save_{idx}"):
                bio = io.BytesIO()
                img.save(bio, format="JPEG")
                thumb_b64 = base64.b64encode(bio.getvalue()).decode('utf-8')
                if "records" not in st.session_state:
                    st.session_state["records"] = []
                st.session_state["records"].append({"name": name, "thumb_b64": thumb_b64, "description": desc_area})
                st.success(f"Baris {idx+1} tersimpan.")

# show saved records
if "records" in st.session_state and st.session_state["records"]:
    st.markdown("### Riwayat Tersimpan (sesi ini)")
    for r in st.session_state["records"][-9:]:
        st.markdown(f"**{r['name'] or 'Anak'}**")
        st.image(Image.open(io.BytesIO(base64.b64decode(r['thumb_b64']))), width=200)
        st.text(r['description'])

# Download Word for last 3 saved records
if st.button("Unduh Laporan Word (.docx) — 3 anak per laporan"):
    recs = st.session_state.get("records", [])
    if len(recs) < 3:
        st.warning("Belum tersimpan 3 baris. Simpan tiap baris sebelum unduh.")
    else:
        take = recs[-3:]
        doc = Document()
        doc.add_heading("TAMAN KANAK–KANAK PERTIWI 26–68", level=1)
        doc.add_paragraph("PRUPUK UTARA KECAMATAN MARGASARI\nKABUPATEN TEGAL\nJl. Anggrek No. 01 Rt. 3 Rw. 4 Desa Prupuk Utara Kec. Margasari Kab. Tegal")
        doc.add_heading("DOKUMEN HASIL KARYA", level=2)
        doc.add_paragraph(f"Kelompok: {kelompok}\tHari/Tanggal: {tanggal}\nSemester: {semester}\tGuru Kelas: {teacher}\nTopik: {topic}\tSub Topik: {subtopic}")
        for rec in take:
            doc.add_heading(f"Nama : {rec['name'] or 'Anak'}", level=3)
            imgdata = base64.b64decode(rec['thumb_b64'])
            with tempfile.NamedTemporaryFile(delete=False, suffix=".jpg") as tf:
                tf.write(imgdata); tf.flush()
                doc.add_picture(tf.name, width=Inches(2.8))
            lines = [ln.strip() for ln in rec['description'].splitlines() if ln.strip()]
            for ln in lines:
                if ln.startswith("•"):
                    doc.add_paragraph(ln[1:].strip(), style='List Bullet')
                else:
                    doc.add_paragraph(ln, style='List Bullet')
            doc.add_page_break()
        bio = io.BytesIO()
        doc.save(bio); bio.seek(0)
        st.download_button("Klik untuk unduh Laporan_Hasil_Karya_Anak.docx", data=bio, file_name="Laporan_Hasil_Karya_Anak.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
